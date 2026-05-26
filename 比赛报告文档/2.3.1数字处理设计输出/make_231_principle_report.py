from __future__ import annotations

from pathlib import Path

import matplotlib as mpl
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch, Rectangle
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
FIG_DIR = BASE_DIR / "figures_principle"
DOCX_PATH = BASE_DIR / "2.3.1 数字处理设计_原理阐述版.docx"

mpl.rcParams.update({
    "font.family": "sans-serif",
    "font.sans-serif": ["Microsoft YaHei", "SimHei", "Arial Unicode MS", "DejaVu Sans", "sans-serif"],
    "axes.unicode_minus": False,
    "svg.fonttype": "none",
    "pdf.fonttype": 42,
    "font.size": 9,
})

INK = "#111827"
MUTED = "#4B5563"
LINE = "#374151"
SOFT = "#9CA3AF"
PANEL = "#F7F8FA"
PANEL2 = "#F2F4F7"
HEADER = "#E5E7EB"
WHITE = "#FFFFFF"

RBW_ROWS = [
    ("1 kHz", 13000, 5, 256, 256, 128),
    ("10 kHz", 1300, 5, 256, 256, 128),
    ("30 kHz", 433, 5, 256, 256, 128),
    ("100 kHz", 130, 5, 128, 384, 64),
    ("300 kHz", 43, 4, 128, 384, 64),
    ("1 MHz", 13, 4, 64, 384, 32),
]

FIG_META = [
    ("fig_231_1_overall_flow", "图 2.3.1-1 数字中频处理总体流程。固定中频信号经采样、正交下变频、RBW滤波、功率检测和幅度换算后形成频谱显示数据。"),
    ("fig_231_2_quadrature_ddc", "图 2.3.1-2 数字正交下变频原理。采样后的实数中频分别与本地数字正弦、余弦相乘，得到 I/Q 两路基带信号。"),
    ("fig_231_3_rbw_filter", "图 2.3.1-3 RBW滤波与等效噪声带宽控制。不同分辨率带宽通过抽取和窄带低通滤波实现，在频率分辨率和扫频速度之间折中。"),
    ("fig_231_4_sweep_spectrum", "图 2.3.1-4 扫频成谱原理。通过改变本振频率依次测量不同射频频点的窄带功率，最终拼接得到完整频谱曲线。"),
    ("fig_231_5_power_calibration", "图 2.3.1-5 功率检测与幅度换算。RBW带宽内的 I/Q 均方功率先换算为 ADC 输入功率，再结合前端增益和损耗修正为射频输入端幅度。"),
]


def setup_canvas(width=12.0, height=7.0):
    fig, ax = plt.subplots(figsize=(width, height))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    return fig, ax


def title(ax, main, sub):
    ax.text(0.04, 0.955, main, ha="left", va="top", fontsize=16, color=INK, fontweight="bold")
    ax.text(0.04, 0.905, sub, ha="left", va="top", fontsize=9.5, color=MUTED)


def box(ax, x, y, w, h, text, fc=WHITE, ec=LINE, fs=10, bold=False):
    patch = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.01,rounding_size=0.012",
        linewidth=1.25,
        edgecolor=ec,
        facecolor=fc,
    )
    ax.add_patch(patch)
    ax.text(
        x + w / 2,
        y + h / 2,
        text,
        ha="center",
        va="center",
        fontsize=fs,
        color=INK,
        fontweight="bold" if bold else "normal",
        linespacing=1.22,
    )


def arrow(ax, start, end, color=LINE, lw=1.4, dashed=False, rad=0.0):
    ax.add_patch(
        FancyArrowPatch(
            start,
            end,
            arrowstyle="-|>",
            mutation_scale=12,
            linewidth=lw,
            color=color,
            linestyle=(0, (5, 4)) if dashed else "solid",
            connectionstyle=f"arc3,rad={rad}",
        )
    )


def save(fig, stem):
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    out = FIG_DIR / stem
    fig.savefig(out.with_suffix(".png"), dpi=300, bbox_inches="tight", facecolor="white")
    fig.savefig(out.with_suffix(".svg"), bbox_inches="tight", facecolor="white")
    fig.savefig(out.with_suffix(".pdf"), bbox_inches="tight", facecolor="white")
    plt.close(fig)


def fig_overall():
    fig, ax = setup_canvas()
    title(ax, "数字中频处理总体流程", "从固定中频采样到频谱曲线显示的信号处理链路。")
    labels = [
        "ADC采样\n固定中频",
        "数字正交\n下变频",
        "I/Q基带\n低速等效信号",
        "RBW滤波\n分辨率带宽",
        "功率检测\n均方能量",
        "幅度换算\n输入端dBm",
        "频谱显示\n频率-幅度曲线",
    ]
    x0, y, w, h, gap = 0.045, 0.56, 0.115, 0.12, 0.024
    for i, text in enumerate(labels):
        x = x0 + i * (w + gap)
        box(ax, x, y, w, h, text, WHITE, LINE, 9.2, True)
        if i < len(labels) - 1:
            arrow(ax, (x + w + 0.003, y + h / 2), (x + w + gap - 0.004, y + h / 2))
    ax.add_patch(Rectangle((0.08, 0.30), 0.26, 0.11, facecolor=PANEL, edgecolor=SOFT, linewidth=1.0, linestyle=(0, (5, 4))))
    ax.text(0.10, 0.38, "频率规划", fontsize=8.8, color=MUTED, va="center")
    ax.text(0.21, 0.34, "RF -> 40 MHz 中频", fontsize=10.2, color=INK, ha="center", va="center")
    ax.add_patch(Rectangle((0.39, 0.30), 0.26, 0.11, facecolor=PANEL, edgecolor=SOFT, linewidth=1.0, linestyle=(0, (5, 4))))
    ax.text(0.41, 0.38, "测量带宽", fontsize=8.8, color=MUTED, va="center")
    ax.text(0.52, 0.34, "RBW 决定分辨率和噪声带宽", fontsize=10.2, color=INK, ha="center", va="center")
    ax.add_patch(Rectangle((0.70, 0.30), 0.22, 0.11, facecolor=PANEL, edgecolor=SOFT, linewidth=1.0, linestyle=(0, (5, 4))))
    ax.text(0.72, 0.38, "成谱方式", fontsize=8.8, color=MUTED, va="center")
    ax.text(0.81, 0.34, "逐频点测量并拼接", fontsize=10.2, color=INK, ha="center", va="center")
    save(fig, "fig_231_1_overall_flow")


def fig_ddc():
    fig, ax = setup_canvas()
    title(ax, "数字正交下变频原理", "将固定中频搬移到零频附近，形成便于滤波和功率检测的复基带信号。")
    box(ax, 0.08, 0.55, 0.15, 0.11, "采样中频\nx[n]", WHITE, LINE, 10, True)
    box(ax, 0.34, 0.68, 0.16, 0.09, "cos(2πfIF n/fs)", PANEL, LINE, 9.5)
    box(ax, 0.34, 0.43, 0.16, 0.09, "-sin(2πfIF n/fs)", PANEL, LINE, 9.5)
    box(ax, 0.58, 0.68, 0.14, 0.09, "低通滤波\nI[n]", WHITE, LINE, 9.5, True)
    box(ax, 0.58, 0.43, 0.14, 0.09, "低通滤波\nQ[n]", WHITE, LINE, 9.5, True)
    box(ax, 0.80, 0.55, 0.13, 0.11, "复基带\nz[n]=I[n]+jQ[n]", WHITE, LINE, 9.2, True)
    arrow(ax, (0.23, 0.605), (0.34, 0.725))
    arrow(ax, (0.23, 0.605), (0.34, 0.475))
    arrow(ax, (0.50, 0.725), (0.58, 0.725))
    arrow(ax, (0.50, 0.475), (0.58, 0.475))
    arrow(ax, (0.72, 0.725), (0.80, 0.62))
    arrow(ax, (0.72, 0.475), (0.80, 0.59))
    ax.text(0.10, 0.28, "作用：把 40 MHz 中频信号变换到基带，使后续 RBW 滤波只需围绕零频设计。", fontsize=10.6, color=INK, fontweight="bold")
    ax.text(0.10, 0.22, "优势：I/Q 表示同时保留幅度和相位信息，避免实数混频后正负频率混叠带来的测量误差。", fontsize=9.5, color=MUTED)
    save(fig, "fig_231_2_quadrature_ddc")


def fig_rbw():
    fig, ax = setup_canvas()
    title(ax, "RBW滤波与等效噪声带宽控制", "分辨率带宽由数字窄带滤波器实现，决定频率分辨能力和噪声功率积分范围。")
    labels = ["I/Q基带", "CIC抽取\n降低数据率", "补偿FIR\n形成RBW通带", "跳过瞬态\n取稳定观测窗", "RBW内\n功率样本"]
    xs = [0.07, 0.25, 0.46, 0.67, 0.84]
    for i, (x, label) in enumerate(zip(xs, labels)):
        box(ax, x, 0.62, 0.13, 0.10, label, WHITE, LINE, 8.8, True)
        if i < len(xs) - 1:
            arrow(ax, (x + 0.13, 0.67), (xs[i + 1], 0.67))
    ax.text(0.08, 0.46, "窄 RBW", fontsize=10.8, color=INK, fontweight="bold")
    ax.add_patch(Rectangle((0.20, 0.46), 0.52, 0.026, facecolor="#D1D5DB", edgecolor=LINE, linewidth=0.5))
    ax.text(0.75, 0.473, "更高频率分辨率 / 更长测量时间 / 更低显示噪声", fontsize=8.7, color=MUTED, va="center")
    ax.text(0.08, 0.36, "宽 RBW", fontsize=10.8, color=INK, fontweight="bold")
    ax.add_patch(Rectangle((0.20, 0.36), 0.18, 0.026, facecolor="#E5E7EB", edgecolor=LINE, linewidth=0.5))
    ax.text(0.42, 0.373, "更快扫频 / 较低频率分辨率 / 更高噪声积分", fontsize=8.7, color=MUTED, va="center")
    ax.text(0.08, 0.22, "本设计提供 1 kHz 至 1 MHz 多档 RBW，用于覆盖窄带信号观察和宽范围快速扫描两类需求。", fontsize=10.0, color=INK, fontweight="bold")
    save(fig, "fig_231_3_rbw_filter")


def fig_sweep():
    fig, ax = setup_canvas()
    title(ax, "扫频成谱原理", "通过改变本振频率，使不同 RF 频率依次落入固定中频处理链。")
    top = [("RF频点 f1", 0.08), ("RF频点 f2", 0.28), ("RF频点 f3", 0.48), ("...", 0.68), ("RF频点 fN", 0.82)]
    for label, x in top:
        box(ax, x, 0.66, 0.12, 0.08, label, WHITE, LINE, 9.2, True)
    for i in range(len(top) - 1):
        arrow(ax, (top[i][1] + 0.12, 0.70), (top[i + 1][1], 0.70))
    box(ax, 0.18, 0.42, 0.22, 0.09, "本振调谐\nLO1 = RF + IF1", PANEL, LINE, 9.4, True)
    box(ax, 0.47, 0.42, 0.22, 0.09, "固定中频链路\n40 MHz IF", PANEL, LINE, 9.4, True)
    box(ax, 0.75, 0.42, 0.16, 0.09, "测得功率\nP(fi)", PANEL, LINE, 9.4, True)
    arrow(ax, (0.29, 0.66), (0.29, 0.51))
    arrow(ax, (0.40, 0.465), (0.47, 0.465))
    arrow(ax, (0.69, 0.465), (0.75, 0.465))
    ax.plot([0.13, 0.25, 0.39, 0.52, 0.66, 0.79, 0.90], [0.22, 0.25, 0.20, 0.34, 0.24, 0.38, 0.21], color=LINE, linewidth=1.8)
    ax.scatter([0.13, 0.25, 0.39, 0.52, 0.66, 0.79, 0.90], [0.22, 0.25, 0.20, 0.34, 0.24, 0.38, 0.21], s=24, color=INK)
    ax.text(0.11, 0.14, "频率", fontsize=9.2, color=MUTED)
    ax.text(0.90, 0.14, "幅度", fontsize=9.2, color=MUTED, ha="right")
    ax.text(0.12, 0.08, "逐点测得的 P(fi) 按频率顺序连接，即形成频谱曲线。", fontsize=10.0, color=INK, fontweight="bold")
    save(fig, "fig_231_4_sweep_spectrum")


def fig_power():
    fig, ax = setup_canvas()
    title(ax, "功率检测与幅度换算", "将 RBW 带宽内的复基带能量换算为射频输入端幅度。")
    labels = ["RBW内\nI/Q样本", "均方功率\nP=mean(I²+Q²)", "相对满量程\nP_dBFS", "ADC输入功率\nP_ADC", "射频输入功率\nP_RF"]
    xs = [0.07, 0.26, 0.46, 0.65, 0.82]
    for i, (x, label) in enumerate(zip(xs, labels)):
        box(ax, x, 0.64, 0.13, 0.10, label, WHITE, LINE, 8.8, True)
        if i < len(xs) - 1:
            arrow(ax, (x + 0.13, 0.69), (xs[i + 1], 0.69))
    box(ax, 0.26, 0.42, 0.20, 0.07, "P_dBFS = 10log10(P/PFS)", PANEL, LINE, 9.2)
    box(ax, 0.55, 0.42, 0.28, 0.07, "P_RF = P_ADC + 前端校正量", PANEL, LINE, 9.2)
    arrow(ax, (0.36, 0.64), (0.36, 0.49), dashed=True)
    arrow(ax, (0.73, 0.64), (0.69, 0.49), dashed=True)
    box(ax, 0.18, 0.22, 0.64, 0.08, "前端校正量综合考虑衰减器、低噪声放大器、可变增益放大器、混频链路损耗和频率响应起伏。", WHITE, LINE, 9.6, True)
    save(fig, "fig_231_5_power_calibration")


def fig_table():
    fig, ax = setup_canvas(width=11.8, height=6.4)
    title(ax, "RBW 档位参数", "不同 RBW 对应不同抽取比和观测长度，用于平衡分辨率与扫频速度。")
    columns = ["RBW", "CIC R", "CIC N", "FIR taps", "Observe pts", "Skip pts", "Accum target"]
    rows = [[rbw, str(r), str(n), str(taps), str(obs), str(skip), str(obs + skip + taps)] for rbw, r, n, taps, obs, skip in RBW_ROWS]
    table = ax.table(cellText=rows, colLabels=columns, cellLoc="center", colLoc="center", bbox=[0.06, 0.20, 0.88, 0.62])
    table.auto_set_font_size(False)
    table.set_fontsize(9.2)
    for (row, col), cell in table.get_celld().items():
        cell.set_edgecolor(SOFT)
        cell.set_linewidth(0.7)
        if row == 0:
            cell.set_facecolor(HEADER)
            cell.get_text().set_color(INK)
            cell.get_text().set_fontweight("bold")
        else:
            cell.set_facecolor("#FFFFFF" if row % 2 else PANEL)
            if col == 0:
                cell.get_text().set_fontweight("bold")
    ax.text(0.06, 0.12, "注：Accum target = Observe points + Skip points + FIR taps；更窄 RBW 通常需要更大的抽取比和更长观测时间。", fontsize=8.7, color=MUTED)
    save(fig, "table_231_1_rbw_params")


def generate_figures():
    fig_overall()
    fig_ddc()
    fig_rbw()
    fig_sweep()
    fig_power()
    fig_table()


def set_doc_defaults(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.color.rgb = RGBColor(17, 24, 39)
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 3"].font.size = Pt(11.5)


def para(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(75, 85, 99)


def add_fig(doc: Document, stem: str, cap: str, width_cm=15.6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG_DIR / f"{stem}.png"), width=Cm(width_cm))
    caption(doc, cap)


def add_table(doc: Document):
    doc.add_heading("表 2.3.1-1 RBW 档位参数", level=3)
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["RBW", "CIC R", "CIC N", "FIR taps", "Observe pts", "Skip pts", "Accum target"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for rbw, r, n, taps, obs, skip in RBW_ROWS:
        vals = [rbw, str(r), str(n), str(taps), str(obs), str(skip), str(obs + skip + taps)]
        cells = table.add_row().cells
        for i, val in enumerate(vals):
            cells[i].text = val
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    caption(doc, "表 2.3.1-1 RBW 档位参数。不同档位通过抽取比、滤波器长度和观测点数的组合实现不同分辨率带宽。")


def generate_word():
    doc = Document()
    set_doc_defaults(doc)
    doc.add_heading("2.3.1 数字处理设计", level=1)
    para(doc, "本仪器采用超外差接收结构，将宽频段射频输入变换到固定中频后再进行数字化处理。数字处理部分的核心目标，是把 ADC 采集到的中频采样序列转换为频谱仪可显示的“频率-幅度”数据。为实现这一目标，系统采用数字正交下变频、分辨率带宽滤波、窄带功率检测和幅度换算等步骤，构成完整的数字中频处理链路。该链路既利用了固定中频结构便于数字处理的特点，又通过扫频方式覆盖 1 MHz-1.5 GHz 的输入范围。")
    para(doc, "从频谱仪测量原理看，显示曲线上的每一个频点并不是一个孤立的采样值，而是该频率附近一定带宽内信号能量的测量结果。这个带宽由分辨率带宽 RBW 决定。RBW 越窄，越容易区分相邻频率分量，显示噪声也更低，但需要更长的观测时间；RBW 越宽，扫频速度更快，但频率分辨能力下降。因此，数字处理设计的重点不只是“采样并显示”，而是要围绕 RBW 建立一套可定量、可调节的窄带功率测量流程。总体流程如图 2.3.1-1 所示。")
    add_fig(doc, *FIG_META[0])

    doc.add_heading("1 数字中频处理总体流程", level=2)
    para(doc, "射频前端完成放大、衰减、滤波和两级混频后，将被测信号搬移到固定的第二中频。固定中频的好处在于，后续数字处理不必随输入射频频率改变而重新设计，只需围绕一个固定中频频率进行数字下变频和滤波。ADC 以高速采样率对该中频信号进行离散化，得到实数采样序列 x[n]。随后数字处理链路首先把 x[n] 搬移到零频附近，形成 I/Q 复基带信号，再根据用户选择的 RBW 对基带信号进行窄带滤波。")
    para(doc, "完成 RBW 滤波后，系统对滤波输出的 I/Q 数据进行能量统计，得到当前频点在该分辨率带宽内的平均功率。该功率先以相对 ADC 满量程的形式表示，再结合 ADC 输入参考、前端增益、衰减和频率响应修正，换算成射频输入端功率。最后，扫频过程中得到的一系列频率点和对应幅度值按频率顺序连接，即形成频谱曲线。")

    doc.add_heading("2 中频采样与数字正交下变频", level=2)
    para(doc, "ADC 输出的是实数中频采样序列，其频谱位于正负中频附近。若直接对该实数序列进行窄带测量，滤波器中心频率需要放在中频位置，且正负频谱成分容易相互影响。数字正交下变频通过在数字域中构造与中频同频率的正弦和余弦本振，将中频信号搬移到零频附近，得到同相分量 I[n] 和正交分量 Q[n]。其基本关系可写为：I[n] = x[n]cos(2πfIFn/fs)，Q[n] = -x[n]sin(2πfIFn/fs)。")
    para(doc, "经过正交混频后，原本位于中频处的信号被搬移到基带，后续只需要围绕零频进行低通滤波即可提取目标频点附近的能量。I/Q 表示的另一个优点是保留了信号的复数形式，能够避免实数处理中的镜像混叠问题，并使功率计算可以直接采用 I²+Q² 的形式。图 2.3.1-2 给出了数字正交下变频的基本结构。")
    add_fig(doc, *FIG_META[1])

    doc.add_heading("3 分辨率带宽滤波设计", level=2)
    para(doc, "RBW 滤波器相当于频谱仪中的数字中频滤波器，它决定了测量时实际积分的频率范围。对于某一射频频点，经过本振变换和数字下变频后，目标信号落在基带附近；此时用低通滤波器限制通过带宽，就可以得到该频点附近 RBW 范围内的信号分量。滤波器的等效噪声带宽越窄，进入功率检测环节的噪声功率越小，显示噪声底越低，但稳定测量所需时间也越长。")
    para(doc, "为了兼顾计算量和滤波效果，数字处理链路采用“抽取加补偿滤波”的思想。首先使用适合大倍率降采样的抽取滤波结构降低数据率，使后续窄带滤波能够在较低等效采样率下完成；随后使用补偿低通滤波器形成更接近目标 RBW 的通带响应，并削弱抽取滤波带来的通带起伏。由于滤波器在启动时存在瞬态响应，功率统计只应在滤波输出进入稳定区间后进行。图 2.3.1-3 展示了 RBW 滤波与观测窗口之间的关系。")
    add_fig(doc, *FIG_META[2])
    para(doc, "本设计设置了从 1 kHz 到 1 MHz 的多档 RBW，以适应不同测量场景。窄 RBW 适合观察相邻很近的窄带信号和弱小杂散，宽 RBW 适合快速寻找信号位置或观察较宽带宽的信号包络。各档 RBW 对应的抽取比、滤波器长度和观测点数见表 2.3.1-1。")
    add_table(doc)

    doc.add_heading("4 扫频成谱与逐点功率检测", level=2)
    para(doc, "本仪器采用扫频式频谱测量方式。对某一个待测频点，射频前端通过调节本振频率使该频点变换到固定中频；数字处理链路再对这个固定中频进行下变频、RBW 滤波和功率检测，得到该频点的幅度值。随后本振切换到下一个频点，重复同样的测量过程。这样得到的一组 P(f1)、P(f2)、...、P(fN) 按频率顺序排列后，就构成了完整频谱。")
    para(doc, "扫频步进需要与 RBW 相匹配。若步进远大于 RBW，窄带信号可能落在两个测量点之间而被漏检；若步进过小，虽然曲线更细密，但扫频时间会明显增加。因此通常使扫频步进与 RBW 保持同量级，并在分辨率和速度之间折中。对于窄 RBW 测量，应使用较小步进以保证窄峰不被跳过；对于宽 RBW 快速扫描，则可以使用较大步进以提高刷新速度。图 2.3.1-4 给出了扫频成谱的基本原理。")
    add_fig(doc, *FIG_META[3])

    doc.add_heading("5 功率计算与幅度换算", level=2)
    para(doc, "经过 RBW 滤波后，I/Q 基带数据表示的是当前频点附近分辨率带宽内的复信号。其瞬时功率与 I²+Q² 成正比，因此可在有效观测窗口内求均值得到该频点的平均功率：P = mean(I²+Q²)。为了便于与 ADC 满量程和实际输入功率对应，首先将该平均功率换算为 dBFS，即相对于 ADC 满量程的分贝值；再根据 ADC 输入满量程对应功率换算为 ADC 输入端 dBm。")
    para(doc, "实际射频输入端功率还需要考虑前端链路的增益和损耗。输入信号经过低噪声放大器、衰减器、混频器、滤波器和可变增益放大器后才进入 ADC，因此 ADC 端测得的功率并不等于被测端口功率。幅度换算时需要把已知的前端增益、衰减、通道损耗和不同频段的频响误差纳入校正，最终得到射频输入端功率。图 2.3.1-5 展示了从 I/Q 样本功率到输入端幅度的换算关系。")
    add_fig(doc, *FIG_META[4])

    doc.add_heading("6 本设计的特点", level=2)
    para(doc, "综上，数字处理链路围绕固定中频和扫频测量两条主线展开。固定中频使数字正交下变频和 RBW 滤波器设计更稳定，扫频方式使仪器能够在较宽射频范围内逐点获得窄带功率。RBW 滤波则把“频谱显示”转化为严格的带宽受限功率测量，使不同档位下的分辨率、噪声底和扫速具有明确物理意义。")
    para(doc, "与直接显示宽带采样频谱相比，该方案更接近传统频谱仪的测量方式，能够在有限采样率和有限处理资源下实现较宽频段覆盖。通过调节本振频率选择待测 RF 频点，通过数字正交下变频把测量集中到基带，通过 RBW 滤波控制分辨率带宽，通过平均功率检测提高幅度稳定性，再通过幅度校正得到可读的输入端功率。这样的处理链路兼顾了低成本实现、测量可解释性和后续标定扩展能力。")

    doc.save(DOCX_PATH)


def main():
    generate_figures()
    generate_word()
    print(f"Generated figures: {FIG_DIR}")
    print(f"Generated Word: {DOCX_PATH}")


if __name__ == "__main__":
    main()
